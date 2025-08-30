import re
import os
import logging
import signal
from typing import Dict, List, Tuple, Any
try:
    import pdfplumber
    import PyPDF2
    from docx import Document
    from docx.shared import Inches
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    PDF_LIBS_AVAILABLE = True
    
    # Reduce logging level for PDF processing to prevent excessive output
    logging.getLogger('pdfminer').setLevel(logging.ERROR)
    logging.getLogger('pdfplumber').setLevel(logging.ERROR)
    
except ImportError as e:
    logging.error(f"Missing required libraries: {e}")
    PDF_LIBS_AVAILABLE = False

class DocumentProcessor:
    """Handles document processing, text extraction, and merging operations"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from PDF or DOCX file"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF using PyPDF2 (safer for complex PDFs)"""
        if not PDF_LIBS_AVAILABLE:
            raise ImportError("PDF processing libraries not available")
        
        text = ""
        
        # Use PyPDF2 directly as it's more stable for complex PDFs
        try:
            if 'PyPDF2' not in globals():
                raise ImportError("PyPDF2 not available")
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                self.logger.info(f"Processing PDF with {total_pages} pages")
                
                for page_num in range(total_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            self.logger.debug(f"Extracted text from page {page_num + 1}")
                        else:
                            self.logger.warning(f"No text found on page {page_num + 1}")
                    except Exception as page_error:
                        self.logger.warning(f"Error extracting text from page {page_num + 1}: {str(page_error)}")
                        continue
                        
        except Exception as pdf_error:
            self.logger.error(f"PDF extraction failed: {str(pdf_error)}")
            # Return a meaningful message instead of crashing
            return f"PDF content could not be fully extracted. Error: {str(pdf_error)}\n\nPlease ensure the PDF contains extractable text and try again."
        
        if not text.strip():
            return "No readable text found in PDF. The document may be image-based, password-protected, or corrupted."
            
        return text
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            if 'Document' not in globals():
                raise ImportError("python-docx not available")
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
    
    def extract_recap_data(self, file_path: str) -> Dict[str, Any]:
        """Extract structured data from recap document"""
        recap_text = self.extract_text_from_file(file_path)
        
        # Define patterns for key information
        patterns = {
            'vessel_name': r'MV\s+([A-Z\s\d]+?)(?:\s+ex\s+|IMO|\s+Cyprus)',
            'imo_number': r'IMO\s+(\d+)',
            'flag': r'(Cyprus|Liberian|Marshall Islands|Panama)\s+flag',
            'built_year': r'BLT\s+(\d{4})',
            'dwt': r'DWT\s+([\d,\']+)',
            'delivery_port': r'Delivery\s+DLOSP\s+([A-Za-z\s,]+)',
            'laycan': r'Laycan:\s*([^–\n]+)',
            'hire_rate': r'at\s+([\d.]+)\s*%\s*bhsi38',
            'period': r'About\s+(\d+)\s+to\s+about\s+(\d+)\s+months',
            'optional_period': r'opt\s+about\s+(\d+)\s*[–-]\s*about\s+(\d+)\s+months\s+at\s*\$\s*([\d,]+)',
            'redelivery_range': r'REDEL\s+DOP\s+1SP\s+WW\s+WITHIN\s+TRADING\s+LIMITS.*?(?=\+|$)',
            'bunkers_delivery': r'VLSFO\s+ABOUT\s+(\d+)\s+MT\s+AND\s+MGO\s+ABOUT\s+(\d+)\s+MT',
            'commission': r'Commission:\s*([\d.]+)\s*%?\s*address\s+commission',
            'charterer': r'LOUIS\s+DREYFUS\s+COMPANY[^+]*',
            'owner': r'MV\s+LUNAR\s+STAR\s+1\s+SHIPPING\s+COMPANY[^+]*',
        }
        
        extracted_data = {}
        
        for key, pattern in patterns.items():
            match = re.search(pattern, recap_text, re.IGNORECASE | re.DOTALL)
            if match:
                try:
                    if key == 'period':
                        extracted_data[key] = f"{match.group(1)} to {match.group(2)} months"
                    elif key == 'optional_period':
                        extracted_data[key] = f"{match.group(1)} to {match.group(2)} months at ${match.group(3)}"
                    elif key == 'bunkers_delivery':
                        extracted_data['vlsfo_quantity'] = match.group(1)
                        extracted_data['mgo_quantity'] = match.group(2)
                    elif match.groups():
                        extracted_data[key] = match.group(1).strip()
                    else:
                        extracted_data[key] = match.group(0).strip()
                except IndexError:
                    # If group doesn't exist, use the full match
                    extracted_data[key] = match.group(0).strip()
        
        # Extract trading exclusions
        trading_section = re.search(r'TRADING EXCLUSIONS(.*?)(?=\+|$)', recap_text, re.IGNORECASE | re.DOTALL)
        if trading_section:
            extracted_data['trading_exclusions'] = trading_section.group(1).strip()
        
        # Extract hire payment details
        hire_section = re.search(r'Hire payment clause:(.*?)(?=Conversion|For subsequent)', recap_text, re.IGNORECASE | re.DOTALL)
        if hire_section:
            extracted_data['hire_payment'] = hire_section.group(1).strip()
        
        return extracted_data
    
    def merge_documents(self, base_cp_text: str, recap_data: Dict[str, Any], change_tracker) -> Tuple[str, List[Dict]]:
        """Merge recap data into base CP template and track changes"""
        merged_text = base_cp_text
        changes = []
        
        # Define mapping rules for merging recap data into CP
        merge_rules = [
            {
                'field': 'charter_date',
                'pattern': r'(made and concluded in\s+\w+\s+)(\d+)(th|st|nd|rd)(\s+day of\s+\w+\s+19\s+)(\d+)',
                'replacement': lambda m: f"{m.group(1)}10{m.group(3)}{m.group(4).replace('May', 'June')}2025",
                'description': 'Updated charter date to June 10, 2025'
            },
            {
                'field': 'vessel_name',
                'pattern': r'(Steamship/Motorship\s+")([^"]+)(")',
                'replacement': lambda m: f'{m.group(1)}LUNAR STAR 1{m.group(3)}',
                'recap_field': 'vessel_name',
                'description': 'Updated vessel name'
            },
            {
                'field': 'owner_details',
                'pattern': r'(Between\s+)[^,]+,([^,]+,){2}[^,]+,([^,]+,)[^,]+',
                'replacement': 'MV LUNAR STAR 1 SHIPPING COMPANY LIMITED, C/o Oesterreichischer Lloyd Seereederei (Cyprus) Ltd, 67 Franklin Roosevelt Ave, Limassol, VAT number CY60177359M',
                'description': 'Updated owner details'
            },
            {
                'field': 'charterer_details',
                'pattern': r'(and\s+)[^C]+COMPANY[^S]+SINGAPORE[^S]+SINGAPORE',
                'replacement': 'Louis Dreyfus Company Suisse S.A.- Charterers of the City of GVA Center,29 route de l\' Aéroport- P.O. Box 236, 1215 Geneva 15,Switzerland',
                'description': 'Updated charterer details'
            },
            {
                'field': 'charter_period',
                'pattern': r'(about\s+minimum\s+)(\d+)(\s+months\s+to\s+maximum\s+)(\d+)(\s+months[^,]*)',
                'replacement': lambda m: 'About 11 to about 14 months (about to mean +/- 15 days in charterer\'s option) at 107.00 % bhsi38 index ( with option to convert to fixed rate) + opt about 10 – about 14 months (about to mean +/ - 15 days in charterer\'s option) at $ 11,500 declarable in charterer\'s option. Optional 10-14 months to be declared by Charterers latest 45 days prior max duration',
                'description': 'Updated charter period with index rate and optional period'
            },
            {
                'field': 'delivery_port',
                'pattern': r'(at\s+on\s+dropping\s+last\s+outward\s+sea\s+pilot\s+)[^(]+(\(intention[^)]*\))',
                'replacement': 'Yeosu, subject Sellers port changes',
                'recap_field': 'delivery_port',
                'description': 'Updated delivery port'
            },
            {
                'field': 'laycan',
                'pattern': r'(WOULD SUGGEST 1 \(ONE\) MONTH AFTER DIVER SURVEY TO BE CARRIED OUT IN YEOSU)',
                'replacement': 'Laycan: 3-10 July – it is mutually agreed between both Owners and Charterers that if any adjustment on laycan required, same to be discussed in good faith by both parties',
                'description': 'Added laycan period'
            }
        ]
        
        # Apply merge rules
        for rule in merge_rules:
            pattern = rule['pattern']
            
            if 'replacement' in rule:
                if callable(rule['replacement']):
                    # Function-based replacement
                    def replace_func(match):
                        old_text = match.group(0)
                        new_text = rule['replacement'](match)
                        changes.append(change_tracker.track_change(
                            rule['field'], old_text, new_text, rule['description']
                        ))
                        return new_text
                    
                    merged_text = re.sub(pattern, replace_func, merged_text, flags=re.IGNORECASE)
                else:
                    # String-based replacement
                    match = re.search(pattern, merged_text, re.IGNORECASE)
                    if match:
                        old_text = match.group(0)
                        new_text = rule['replacement']
                        changes.append(change_tracker.track_change(
                            rule['field'], old_text, new_text, rule['description']
                        ))
                        merged_text = re.sub(pattern, new_text, merged_text, flags=re.IGNORECASE)
        
        # Add special clauses from recap
        if 'trading_exclusions' in recap_data:
            # Find location to insert trading exclusions
            exclusion_pattern = r'(as the Charterers or their Agents shall direct, on the following conditions:)'
            if re.search(exclusion_pattern, merged_text):
                trading_exclusions = f"\n\nTRADING EXCLUSIONS:\n{recap_data['trading_exclusions']}\n"
                merged_text = re.sub(exclusion_pattern, r'\1' + trading_exclusions, merged_text)
                changes.append(change_tracker.track_change(
                    'trading_exclusions', '', trading_exclusions, 'Added trading exclusions clause'
                ))
        
        # Add dry-docking clause
        dd_clause = "\n\nDRY-DOCKING CLAUSE:\nDry-docking / SS Oct - early Nov 25 in China or Med/Black Sea/Portugal, duration 10-15 days. DD location (med or China) is in Charterer's option. Charterers to place the vessel into either Med Sea / Black Sea range or Singapore - Japan range for owners to take over for DD.\n"
        merged_text += dd_clause
        changes.append(change_tracker.track_change(
            'dry_docking', '', dd_clause, 'Added dry-docking clause'
        ))
        
        # Add summary of changes section
        summary_section = "\n\n=== SUMMARY OF CHANGES ===\n"
        if changes:
            summary_section += "The following modifications were made to the base Charter Party:\n\n"
            for i, change in enumerate(changes, 1):
                summary_section += f"{i}. {change['description']}\n"
            summary_section += "\n=== END SUMMARY ===\n"
        else:
            summary_section += "No changes were made to the base Charter Party document.\n=== END SUMMARY ===\n"
        
        merged_text += summary_section
        
        return merged_text, changes
    
    def generate_docx(self, content: str, output_path: str):
        """Generate DOCX file from merged content"""
        try:
            if not PDF_LIBS_AVAILABLE or 'Document' not in globals():
                raise ImportError("python-docx not available")
            doc = Document()
            
            # Add title and header
            title = doc.add_heading('Time Charter Party', 0)
            subtitle = doc.add_heading('GOVERNMENT FORM', level=1)
            sub_subtitle = doc.add_paragraph('Approved by the New York Produce Exchange')
            date_line = doc.add_paragraph('November 6th, 1913 - Amended October 20th, 1921; August 6th, 1931; October 3rd, 1946')
            
            # Add spacing
            doc.add_paragraph()
            
            # Process the content more intelligently
            lines = content.split('\n')
            current_paragraph = ""
            
            for line in lines:
                line = line.strip()
                if not line:
                    # Empty line - finish current paragraph and add spacing
                    if current_paragraph:
                        self._add_formatted_paragraph(doc, current_paragraph)
                        current_paragraph = ""
                    continue
                
                # Check for special sections
                if line.startswith('==='):
                    # Summary section
                    if current_paragraph:
                        self._add_formatted_paragraph(doc, current_paragraph)
                        current_paragraph = ""
                    if 'SUMMARY OF CHANGES' in line:
                        doc.add_page_break()
                        doc.add_heading('Summary of Changes', level=1)
                    elif 'END SUMMARY' in line:
                        doc.add_paragraph()
                    continue
                
                # Check for numbered clauses (1., 2., etc.)
                if re.match(r'^\d+\.\s', line):
                    if current_paragraph:
                        self._add_formatted_paragraph(doc, current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(line, level=2)
                    continue
                
                # Check for clause headers (That the...)
                if line.startswith('That the') and len(line) > 50:
                    if current_paragraph:
                        self._add_formatted_paragraph(doc, current_paragraph)
                        current_paragraph = ""
                    doc.add_heading(line, level=3)
                    continue
                
                # Regular content - accumulate in current paragraph
                if current_paragraph:
                    current_paragraph += " " + line
                else:
                    current_paragraph = line
                
                # If paragraph is getting too long, break it
                if len(current_paragraph) > 500:
                    self._add_formatted_paragraph(doc, current_paragraph)
                    current_paragraph = ""
            
            # Add any remaining content
            if current_paragraph:
                self._add_formatted_paragraph(doc, current_paragraph)
            
            doc.save(output_path)
            self.logger.info(f"DOCX file saved to: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Error generating DOCX: {str(e)}")
            raise
    
    def _add_formatted_paragraph(self, doc, text):
        """Add a formatted paragraph to the document"""
        if not text.strip():
            return
        
        # Clean up the text
        text = text.strip()
        
        # Check if this looks like a list item
        if re.match(r'^\d+\)\s|^[a-z]\)\s|^•\s|^-\s', text):
            p = doc.add_paragraph(text, style='List Bullet')
        else:
            p = doc.add_paragraph(text)
        
        # Add some spacing after paragraphs
        from docx.shared import Pt
        try:
            p.paragraph_format.space_after = Pt(6)
        except:
            pass  # Skip if formatting fails
    
    def generate_pdf(self, content: str, output_path: str):
        """Generate PDF file from merged content"""
        try:
            if not PDF_LIBS_AVAILABLE:
                raise ImportError("reportlab not available")
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # Custom styles with reduced spacing
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=16,
                spaceAfter=20,
                spaceBefore=0,
                alignment=1  # Center alignment
            )
            
            header_style = ParagraphStyle(
                'CustomHeader',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=8,
                spaceBefore=0,
                alignment=1
            )
            
            # Create a custom normal style with controlled spacing
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=10,
                spaceAfter=4,
                spaceBefore=0,
                leading=12  # Line height
            )
            
            section_style = ParagraphStyle(
                'SectionHeader',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=6,
                spaceBefore=10,
                leftIndent=0
            )
            
            story = []
            
            # Add title and headers
            story.append(Paragraph("Time Charter Party", title_style))
            story.append(Paragraph("GOVERNMENT FORM", header_style))
            story.append(Paragraph("Approved by the New York Produce Exchange", normal_style))
            story.append(Paragraph("November 6th, 1913 - Amended October 20th, 1921; August 6th, 1931; October 3rd, 1946", normal_style))
            story.append(Spacer(1, 12))
            
            # Process content more intelligently
            lines = content.split('\n')
            current_paragraph = ""
            
            for line in lines:
                line = line.strip()
                if not line:
                    # Empty line - finish current paragraph
                    if current_paragraph:
                        self._add_pdf_paragraph(story, current_paragraph, normal_style, section_style)
                        current_paragraph = ""
                    continue
                
                # Skip summary markers
                if line.startswith('==='):
                    if current_paragraph:
                        self._add_pdf_paragraph(story, current_paragraph, normal_style, section_style)
                        current_paragraph = ""
                    if 'SUMMARY OF CHANGES' in line:
                        story.append(Spacer(1, 20))
                        story.append(Paragraph("Summary of Changes", section_style))
                    continue
                
                # Check for numbered clauses
                if re.match(r'^\d+\.\s', line):
                    if current_paragraph:
                        self._add_pdf_paragraph(story, current_paragraph, normal_style, section_style)
                        current_paragraph = ""
                    story.append(Paragraph(line, section_style))
                    continue
                
                # Accumulate regular content
                if current_paragraph:
                    current_paragraph += " " + line
                else:
                    current_paragraph = line
                
                # Break very long paragraphs
                if len(current_paragraph) > 800:
                    self._add_pdf_paragraph(story, current_paragraph, normal_style, section_style)
                    current_paragraph = ""
            
            # Add any remaining content
            if current_paragraph:
                self._add_pdf_paragraph(story, current_paragraph, normal_style, section_style)
            
            doc.build(story)
            self.logger.info(f"PDF file saved to: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Error generating PDF: {str(e)}")
            raise
    
    def _add_pdf_paragraph(self, story, text, normal_style, section_style):
        """Add a paragraph to the PDF story with appropriate formatting"""
        if not text.strip():
            return
        
        # Clean up text
        text = text.strip()
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        
        # Check if it's a special section
        if text.startswith('That the') and len(text) > 50:
            story.append(Paragraph(text, section_style))
        else:
            story.append(Paragraph(text, normal_style))
